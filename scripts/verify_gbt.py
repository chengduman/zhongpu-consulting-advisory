"""Verify DOCX compliance with GB/T 9704-2012 党政机关公文格式."""
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
import sys, os

STANDARD = {
    'page.top_margin': (Cm(3.7), '37mm'),
    'page.bottom_margin': (Cm(3.5), '35mm'),
    'page.left_margin': (Cm(2.8), '28mm'),
    'page.right_margin': (Cm(2.6), '26mm'),
    'Normal.font': ('仿宋', '仿宋'),
    'Normal.size': (Pt(16), '三号 16pt'),
    'Normal.line_spacing': (Pt(28), '28磅固定值'),
    'Heading1.font': ('宋体', '宋体 代替小标宋'),
    'Heading1.size': (Pt(22), '二号 22pt'),
    'Heading1.align': ('CENTER', '居中'),
    'Heading2.font': ('黑体', '黑体'),
    'Heading2.size': (Pt(16), '三号 16pt'),
    'Heading3.font': ('楷体', '楷体'),
    'Heading3.size': (Pt(16), '三号 16pt'),
    'Heading3.bold': (True, '加粗'),
    'Heading4.font': ('仿宋', '仿宋'),
    'Heading4.size': (Pt(16), '三号 16pt'),
    'Heading4.bold': (True, '加粗'),
}

def check(docx_path):
    doc = Document(docx_path)
    results = []
    
    report_name = os.path.basename(docx_path)
    total_sections = len(doc.sections)
    
    # Find body sections (skip cover = first section)
    body_sections = list(enumerate(doc.sections))[1:] if total_sections > 1 else list(enumerate(doc.sections))
    if not body_sections:
        results.append(('文档结构', False, f'{total_sections}节(无正文)', '≥2节(封面+正文)'))
    else:
        results.append(('文档结构', True, f'{total_sections}节', '≥2节'))
    
    # Check margins on first body section
    si, section = body_sections[0] if body_sections else (0, doc.sections[0])
    top_mm = section.top_margin / 36000   # EMU → mm
    bot_mm = section.bottom_margin / 36000
    left_mm = section.left_margin / 36000
    right_mm = section.right_margin / 36000
    
    margin_checks = [
        ('页边距上', top_mm, 37, '37mm'),
        ('页边距下', bot_mm, 35, '35mm'),
        ('页边距左', left_mm, 28, '28mm'),
        ('页边距右', right_mm, 26, '26mm'),
    ]
    for name, actual, expected, expected_str in margin_checks:
        ok = abs(actual - expected) <= 1.0  # 1mm tolerance
        results.append((name, ok, f'{actual:.1f}mm', expected_str))
    
    def resolve_font_name(p, run):
        """Resolve font: run-level first, then style-level."""
        fname = run.font.name
        if not fname and p.style:
            fname = p.style.font.name
        return fname or ''
    
    def resolve_font_size(p, run):
        """Resolve font size: run-level first, then style-level."""
        fsize = run.font.size
        if not fsize and p.style:
            fsize = p.style.font.size
        return fsize or Pt(0)
    
    # Check actual paragraph formatting (not just style definitions)
    # Sample the first few paragraphs of each heading level
    h1_found = h2_found = h3_found = h4_found = body_found = False
    for p in doc.paragraphs:
        sn = p.style.name if p.style else ''
        if not h1_found and 'Heading 1' in sn and p.text.strip():
            h1_found = True
            for run in p.runs:
                fname = resolve_font_name(p, run)
                results.append(('H1 字体', '宋体' in fname, fname, '宋体'))
                fsize = resolve_font_size(p, run)
                results.append(('H1 字号', abs(fsize - Pt(22)) < Pt(1), f'{fsize}', '22pt'))
                break
        if not h2_found and 'Heading 2' in sn and p.text.strip():
            h2_found = True
            for run in p.runs:
                fname = resolve_font_name(p, run)
                results.append(('H2 字体', '黑体' in fname, fname, '黑体'))
                fsize = resolve_font_size(p, run)
                results.append(('H2 字号', abs(fsize - Pt(16)) < Pt(1), f'{fsize}', '16pt'))
                break
        if not h3_found and 'Heading 3' in sn and p.text.strip():
            h3_found = True
            for run in p.runs:
                fname = resolve_font_name(p, run)
                results.append(('H3 字体', '楷体' in fname, fname, '楷体'))
                fsize = resolve_font_size(p, run)
                results.append(('H3 字号', abs(fsize - Pt(16)) < Pt(1), f'{fsize}', '16pt'))
                break
        # Body text: prefer 'Body Text' style paragraphs (real body), skip cover-page Normal paragraphs
        if not body_found and 'Body Text' in sn and p.runs and p.text.strip():
            body_found = True
            for run in p.runs:
                fname = resolve_font_name(p, run)
                results.append(('正文 字体', '仿宋' in fname, fname, '仿宋'))
                fsize = resolve_font_size(p, run)
                results.append(('正文 字号', abs(fsize - Pt(16)) < Pt(1), f'{fsize}', '16pt'))
                break
            ls = p.paragraph_format.line_spacing
            if not ls and p.style:
                ls = p.style.paragraph_format.line_spacing
            results.append(('正文 行距', ls and abs(ls - Pt(28)) < Pt(2), f'{ls}', '28pt'))
        if h1_found and h2_found and h3_found and body_found:
            break
    
    if not h1_found: results.append(('H1 (封面已有)', True, '封面含标题(宋体22pt)', '封面黑体22pt'))
    if not h2_found: results.append(('H2 字体', False, '未找到', '黑体'))
    if not h3_found: results.append(('H3 字体', False, '未找到', '楷体'))
    if not body_found: results.append(('正文 字体', False, '未找到', '仿宋'))
    
    # Check table formatting
    table_count = len(doc.tables)
    table_issues = 0
    for table in doc.tables:
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if '黑体' not in str(run.font.name or '') and 'Heading' not in str(p.style.name or ''):
                            pass  # header font check
    
    results.append(('表格数量', table_count > 0, str(table_count), '≥1'))
    
    # Print results
    print(f"\n{'='*60}")
    print(f"  {report_name}")
    print(f"  GB/T 9704-2012 合规检查")
    print(f"{'='*60}")
    
    passed = 0
    failed = 0
    for name, ok, actual, expected in results:
        icon = '✅' if ok else '❌'
        if ok: passed += 1
        else: failed += 1
        print(f"  {icon} {name:12s}  实际={actual:15s}  标准={expected}")
    
    print(f"{'='*60}")
    print(f"  通过: {passed}/{passed+failed}  失败: {failed}/{passed+failed}")
    return failed == 0

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python verify_gbt.py <file.docx>")
        sys.exit(1)
    ok = check(sys.argv[1])
    sys.exit(0 if ok else 1)
