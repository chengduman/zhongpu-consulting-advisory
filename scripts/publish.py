"""Zhongpu Document Publisher — GB/T 9704-2012 compliant."""
import subprocess, os, sys, yaml, tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

PANDOC = r"C:\Users\cheng\AppData\Local\Pandoc\pandoc"
TEMPLATE = r"C:\Users\cheng\zhongpu-consulting-advisory\templates\reference.docx"

def parse_frontmatter(path):
    with open(path, 'r', encoding='utf-8') as f:
        c = f.read()
    if c.startswith('---'):
        parts = c.split('---', 2)
        return (yaml.safe_load(parts[1]), parts[2]) if len(parts) >= 3 else ({}, c)
    return {}, c

def create_cover(meta, out):
    doc = Document()
    s = doc.sections[0]; s.page_width, s.page_height = Cm(21), Cm(29.7)
    for a in ['top_margin','bottom_margin','left_margin','right_margin']:
        setattr(s, a, Cm(0))
    NAVY=RGBColor(0x0B,0x1D,0x3A); GOLD=RGBColor(0xC4,0xA4,0x5A)
    GRAY=RGBColor(0x66,0x66,0x66); LIGHT=RGBColor(0x99,0x99,0x99)

    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(0)
    r=p.add_run(' ');r.font.size=Pt(80)
    p.paragraph_format.element.get_or_add_pPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="0B1D3A" w:val="clear"/>'))

    for txt,font,sz,color,bold in[('ZHONGPU','Arial',11,LIGHT,False),('中 普 咨 询','黑体',26,NAVY,True)]:
        pp=doc.add_paragraph();pp.alignment=WD_ALIGN_PARAGRAPH.CENTER;pp.paragraph_format.space_after=Pt(0)
        rr=pp.add_run(txt);rr.font.name=font;rr.font.size=Pt(sz);rr.font.color.rgb=color;rr.bold=bold
        rr.element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')

    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(20);p.paragraph_format.space_after=Pt(20)
    r=p.add_run('─'*24);r.font.size=Pt(6);r.font.color.rgb=GOLD

    cls=meta.get('classification','内部使用')
    pp=doc.add_paragraph();pp.alignment=WD_ALIGN_PARAGRAPH.CENTER;pp.paragraph_format.space_after=Pt(6)
    rr=pp.add_run(cls);rr.font.name='黑体';rr.font.size=Pt(11);rr.font.color.rgb=GOLD;rr.bold=True
    rr.element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')

    proj=meta.get('project_number','')
    if proj:
        pp=doc.add_paragraph();pp.alignment=WD_ALIGN_PARAGRAPH.CENTER;pp.paragraph_format.space_after=Pt(32)
        rr=pp.add_run(proj);rr.font.name='Arial';rr.font.size=Pt(9);rr.font.color.rgb=LIGHT

    title=meta.get('title','').strip('"').strip("'")
    pp=doc.add_paragraph();pp.alignment=WD_ALIGN_PARAGRAPH.CENTER;pp.paragraph_format.space_after=Pt(8)
    rr=pp.add_run(title);rr.font.name='黑体';rr.font.size=Pt(22);rr.bold=True;rr.font.color.rgb=NAVY
    rr.element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')

    sub=meta.get('subtitle','').strip('"').strip("'")
    if sub:
        pp=doc.add_paragraph();pp.alignment=WD_ALIGN_PARAGRAPH.CENTER;pp.paragraph_format.space_after=Pt(50)
        rr=pp.add_run(sub);rr.font.name='等线';rr.font.size=Pt(10.5);rr.font.color.rgb=GRAY
        rr.element.rPr.rFonts.set(qn('w:eastAsia'),'等线')

    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(24)
    r=p.add_run('─'*20);r.font.size=Pt(4);r.font.color.rgb=RGBColor(0xCC,0xCC,0xCC)

    author=meta.get('author','中普咨询').strip('"').strip("'")
    date=meta.get('date','2026年6月')
    info=[]
    if meta.get('department'):info.append(('部门',meta['department']))
    info+=[('编制',author),('日期',date)]
    for label,value in info:
        pp=doc.add_paragraph();pp.alignment=WD_ALIGN_PARAGRAPH.CENTER;pp.paragraph_format.space_after=Pt(4)
        r1=pp.add_run(f'{label}  ');r1.font.name='等线';r1.font.size=Pt(9);r1.font.color.rgb=LIGHT
        r1.element.rPr.rFonts.set(qn('w:eastAsia'),'等线')
        r2=pp.add_run(value);r2.font.name='等线';r2.font.size=Pt(9);r2.font.color.rgb=GRAY
        r2.element.rPr.rFonts.set(qn('w:eastAsia'),'等线')

    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(60)
    r=p.add_run(' ');r.font.size=Pt(20)
    p.paragraph_format.element.get_or_add_pPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="0B1D3A" w:val="clear"/>'))
    doc.save(out)

def main():
    if len(sys.argv) < 2:
        print("Usage: python publish.py <report.md>"); sys.exit(1)
    src = sys.argv[1]; base = os.path.splitext(src)[0]
    meta, body = parse_frontmatter(src)
    title = meta.get('title', os.path.basename(src)).strip('"').strip("'")

    print(f"[1/4] 封面: {title}")
    cover_path = os.path.join(tempfile.gettempdir(), 'zp_cover.docx')
    create_cover(meta, cover_path)

    print("[2/4] DOCX...")
    body_docx = os.path.join(tempfile.gettempdir(), 'zp_body.docx')
    with tempfile.NamedTemporaryFile(suffix='.md', mode='w', encoding='utf-8', delete=False) as f:
        f.write(body); tmp_md = f.name
    subprocess.run([PANDOC, tmp_md, '-o', body_docx, '--reference-doc', TEMPLATE,
        '--from','markdown+pipe_tables','--to','docx','--metadata','lang=zh-CN'], check=True)
    os.unlink(tmp_md)

    # Merge: load pandoc output, prepend cover page (AFTER TOC)
    body_doc = Document(body_docx)
    cover_doc = Document(cover_path)
    
    # Find the 目录 (TOC) heading in body
    toc_index = None
    for i, p in enumerate(body_doc.paragraphs):
        if p.text.strip() == '目录' and ('Heading' in (p.style.name or '')):
            toc_index = i
            break
    
    # Insert cover elements BEFORE the body, but AFTER the TOC
    # This keeps TOC bookmark positions correct
    body_element = body_doc.element.body
    insert_point = body_element[toc_index] if toc_index else body_element[0]
    for el in reversed(list(cover_doc.element.body)):
        body_element.insert(list(body_element).index(insert_point), el)
    
    # The first section (now cover) gets 0 margins
    # Pandoc's sections already have GB/T margins from reference.docx
    # Just ensure cover section has 0 margins
    for i, sec in enumerate(body_doc.sections):
        if i == 0:
            sec.top_margin = Cm(0); sec.bottom_margin = Cm(0)
            sec.left_margin = Cm(0); sec.right_margin = Cm(0)

    # Add headers + paragraph formatting to body sections
    for i, sec in enumerate(body_doc.sections):
        if i == 0: continue
        hdr = sec.header; hdr.is_linked_to_previous = False
        hp = hdr.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hp.add_run(f'中普咨询 | {meta.get("project_number","")} | 机密')
        r.font.name = '仿宋'; r.font.size = Pt(7.5); r.font.color.rgb = RGBColor(0xAA,0xAA,0xAA)
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    # GB/T 9704-2012: body paragraphs first-line indent 2 chars
    from docx.shared import Cm as CmShared
    for p in body_doc.paragraphs:
        sn = p.style.name or ''
        if 'Body' in sn and p.text.strip() and not p.text.startswith('|'):
            p.paragraph_format.first_line_indent = Pt(32)  # 2 chars at 16pt

    # Table formatting: borders + header shading + optimized layout
    for table in body_doc.tables:
        tbl = table._tbl; tblPr = tbl.tblPr or parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        # Full borders
        borders = parse_xml(f'<w:tblBorders {nsdecls("w")}>'
            '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '</w:tblBorders>')
        for ex in tblPr.findall(qn('w:tblBorders')): tblPr.remove(ex)
        tblPr.append(borders)
        # Auto-fit to window
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
        for ex in tblPr.findall(qn('w:tblW')): tblPr.remove(ex)
        tblPr.append(tblW)
        # Header row: repeat on each page + gray background
        if table.rows:
            tr = table.rows[0]._tr
            trPr = tr.find(qn('w:trPr'))
            if trPr is None:
                trPr = parse_xml(f'<w:trPr {nsdecls("w")}/>')
                tr.insert(0, trPr)
            # Repeat header
            hdr_flag = parse_xml(f'<w:tblHeader {nsdecls("w")}/>')
            for ex in trPr.findall(qn('w:tblHeader')): trPr.remove(ex)
            trPr.append(hdr_flag)
            # Shading + font
            for cell in table.rows[0].cells:
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8E8E8" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shd)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = '黑体'
                        run.font.size = Pt(10.5)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0,0,0)
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        # Body rows: smaller font
        for row in table.rows[1:]:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = '仿宋'
                        run.font.size = Pt(10.5)
                        run.font.color.rgb = RGBColor(0,0,0)
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    docx_path = base + '.docx'
    body_doc.save(docx_path)
    print(f"      {docx_path} ({os.path.getsize(docx_path):,} bytes)")

    print("[3/4] HTML5...")
    html_path = base + '.html'
    css_path = src.replace('.md','.css')
    with open(css_path,'w',encoding='utf-8') as f:
        f.write('''@page{size:A4;margin:2.54cm 3.17cm}
body{font-family:"DengXian","Microsoft YaHei",sans-serif;font-size:10.5pt;line-height:1.8;color:#222}
h1{font-family:"SimHei",sans-serif;font-size:22pt;text-align:center}
h2{font-family:"SimHei",sans-serif;font-size:15pt}
h3{font-family:"SimHei",sans-serif;font-size:12pt;color:#333}
table{border-collapse:collapse;width:100%;font-size:9pt}
th{background:#2d2d2d;color:#fff;padding:5pt 8pt}
td{border:1px solid #ccc;padding:4pt 8pt}
tr:nth-child(even) td{background:#fafafa}
blockquote{font-family:"KaiTi",serif;font-size:9pt;color:#666}''')
    subprocess.run([PANDOC, src, '-o', html_path, '--standalone', '--css', css_path,
        '--from','markdown+pipe_tables','--to','html5','--metadata','lang=zh-CN'], check=True)
    os.remove(css_path)
    print(f"      {html_path} ({os.path.getsize(html_path):,} bytes)")

    for t in [cover_path, body_docx]:
        try: os.remove(t)
        except: pass
    print("[4/4] ✅ 出版完成")

if __name__ == '__main__': main()
